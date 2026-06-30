from __future__ import annotations

from pathlib import Path

from app.plantilla_generator.replacements import ReplacementRule, apply_text_replacements


class WordCleaner:
    def clean(
        self,
        source: str | Path,
        destination: str | Path,
        rules: list[ReplacementRule],
    ) -> tuple[int, list[str]]:
        from docx import Document

        document = Document(str(source))
        total = 0
        values: list[str] = []

        for paragraph in document.paragraphs:
            count, replaced = self._replace_in_paragraph(paragraph, rules)
            total += count
            values.extend(replaced)

        for table in document.tables:
            count, replaced = self._replace_in_table(table, rules)
            total += count
            values.extend(replaced)

        for section in document.sections:
            for container in (section.header, section.footer):
                for paragraph in container.paragraphs:
                    count, replaced = self._replace_in_paragraph(paragraph, rules)
                    total += count
                    values.extend(replaced)
                for table in container.tables:
                    count, replaced = self._replace_in_table(table, rules)
                    total += count
                    values.extend(replaced)

        document.save(str(destination))
        return total, sorted(set(values))

    def _replace_in_table(self, table: object, rules: list[ReplacementRule]) -> tuple[int, list[str]]:
        total = 0
        values: list[str] = []
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count, replaced = self._replace_in_paragraph(paragraph, rules)
                    total += count
                    values.extend(replaced)
                for nested_table in cell.tables:
                    count, replaced = self._replace_in_table(nested_table, rules)
                    total += count
                    values.extend(replaced)
        return total, values

    def _replace_in_paragraph(
        self,
        paragraph: object,
        rules: list[ReplacementRule],
    ) -> tuple[int, list[str]]:
        total = 0
        values: list[str] = []
        for run in paragraph.runs:
            new_text, count, replaced = apply_text_replacements(run.text, rules)
            if count:
                run.text = new_text
                total += count
                values.extend(replaced)
        if total:
            return total, values

        original = paragraph.text
        new_text, count, replaced = apply_text_replacements(original, rules)
        if not count:
            return 0, []
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new_text)
        return count, replaced
